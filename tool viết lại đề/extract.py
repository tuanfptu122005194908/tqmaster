#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
extract.py - Trich xuat hang loat de thi trac nghiem tu anh.

Pipeline 3 buoc (nhanh + chinh xac):
  1. Tesseract OCR  -> doc anh song song (nhieu process cung luc)
  2. Gom tat ca text -> 1 file all_ocr.txt (de debug)
  3. Groq AI        -> batch nhieu cau / 1 API call -> JSON chuan

Cach dung:
    1. Bo toan bo anh de thi (.png/.jpg/.jpeg/.webp) vao thu muc input_images/
    2. Chay: python extract.py
    3. Ket qua: output/questions.docx  va  output/all_raw.json

Yeu cau:
    - Tesseract v5+ da cai san (co trong PATH)
    - pip install -r requirements.txt
"""

import sys
import io

# Fix encoding cho console Windows
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
else:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import json
import re
import time
import subprocess
import threading
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed

from groq import Groq

try:
    from json_repair import repair_json
    HAS_JSON_REPAIR = True
except ImportError:
    HAS_JSON_REPAIR = False

from docx import Document

# ----------------------------------------------------------------------------
# Cau hinh
# ----------------------------------------------------------------------------
INPUT_DIR   = Path(__file__).parent / "input_images"
OUTPUT_DIR  = Path(__file__).parent / "output"
OCR_RAW_DIR = OUTPUT_DIR / "ocr_raw"

import os

GROQ_API_KEYS = [
    os.environ.get("GROQ_API_KEY_1", "placeholder_groq_1"),
    os.environ.get("GROQ_API_KEY_2", "placeholder_groq_2"),
]
GROQ_MODEL       = "llama-3.3-70b-versatile"  # text model, nhanh
MAX_JSON_RETRIES = 3
MAX_OCR_WORKERS  = 8    # Tesseract chay song song (subprocess, khong bi GIL)
MAX_GROQ_WORKERS = 4    # Groq API call song song
BATCH_SIZE       = 5    # So cau moi Groq batch call (5 cau -> 1 API call)

SUPPORTED_EXT = {".png", ".jpg", ".jpeg", ".webp"}

# Tesseract: dung ca tieng Anh + tieng Viet, psm 6 = assume uniform block of text
TESS_LANG = "eng+vie"
TESS_PSM  = "6"

# ----------------------------------------------------------------------------
# Thread-safe print
# ----------------------------------------------------------------------------
_print_lock = threading.Lock()

def safe_print(msg: str):
    with _print_lock:
        print(msg, flush=True)

# ----------------------------------------------------------------------------
# Prompt cho Groq (xu ly BATCH nhieu cau 1 luc)
# ----------------------------------------------------------------------------
BATCH_PROMPT_TEMPLATE = """\
Day la text tho duoc doc bang OCR (Tesseract) tu {n} anh cau hoi trac nghiem.
Moi anh duoc ngan cach boi "=== CAU {idx} | file: {fname} ===" va ket thuc bang "--- END CAU {idx} ---".

Anh goc co watermark lon giua (HAWK, FAWK, HAAWK, FWK, WK, DAWK, DONG DAWK) - chi la watermark, KHONG phai noi dung.

OCR co the bi:
- Watermark xen vao text -> XOA bo hoan toan
- Dap an bi lap (xuat hien 2 lan) -> chi giu trong options
- Ky tu la, dong thua -> sua sach
- Option A/B/C/D co the bi OCR doc sai thu tu hoac thieu -> suy luan dua vao ngu canh

NHIEM VU: Phan tich tung cau, tra ve JSON array voi {n} phan tu, moi phan tu theo schema:
{{
  "id": "<ma de + so cau neu doc duoc, neu khong dung ten file>",
  "source_title": "<dong tieu de goc neu co>",
  "answer": ["<chu cai dap an dung neu co, mang rong [] neu khong ro>"],
  "answer_count_note": "<ghi chu neu co, vi du Choose 3 answers>",
  "question": "<FULL phan than cau hoi. KHONG chua bat ky lua chon A/B/C/D. KHONG chua watermark. Giu xuong dong bang \\n>",
  "options": {{
    "A": "<Noi dung lua chon A, day du, sach>",
    "B": "<Noi dung lua chon B>",
    "C": "<Noi dung lua chon C>",
    "D": "<Noi dung lua chon D, them E/F neu co>"
  }},
  "_source_file": "<ten file anh goc>",
  "ocr_confidence": "<high / low>",
  "uncertain_parts": "<mo ta phan sai/mo, de trong neu high>"
}}

KHONG markdown, KHONG giai thich, chi JSON array thuan tuy bat dau bang [ va ket thuc bang ].

Luu y BAT BUOC:
- Doc DAY DU ca 4 lua chon A, B, C, D. KHONG bo sot bat ky option nao.
- "question" CHI chua phan than de bai, KHONG chua A., B., C., D.
- Watermark (HAWK, FAWK, WK...) -> XOA hoan toan, khong de vao bat cu truong nao.
- Neu mot option bi cat giua chuoi (vi du: "Avoid workplace conflict and") -> su dung ngu canh de HOAN THIEN option do (vi du: "Avoid workplace conflict and maintain professionalism"). Ghi ro vao uncertain_parts.
- Neu option bi mat hoan toan do watermark che -> ghi "[khong doc duoc do watermark]" vao o do va ghi vao uncertain_parts.
- Neu co code (JS, SQL, Python...) thi sao chep nguyen van.


TEXT THO OCR (gom {n} cau):
---
{batch_text}
---
"""

# ----------------------------------------------------------------------------
# Groq client pool (round-robin, thread-safe)
# ----------------------------------------------------------------------------
_groq_clients: dict = {}
_groq_key_idx  = 0
_groq_key_lock = threading.Lock()


def get_groq_client(key: str) -> Groq:
    if key not in _groq_clients:
        _groq_clients[key] = Groq(api_key=key)
    return _groq_clients[key]


def next_key() -> str:
    global _groq_key_idx
    with _groq_key_lock:
        key = GROQ_API_KEYS[_groq_key_idx % len(GROQ_API_KEYS)]
        _groq_key_idx += 1
    return key


# ----------------------------------------------------------------------------
# Stage 1: Tesseract OCR song song (subprocess)
# ----------------------------------------------------------------------------
def ocr_one(image_path: Path) -> tuple:
    """
    Chay Tesseract OCR cho 1 anh, tra ve (path, raw_text).
    - Convert anh sang PNG (in-memory) truoc khi OCR de Tesseract doc duoc moi dinh dang.
    - Dung pipe stdin thay vi file tam -> khong can ghi disk.
    - Chay trong ThreadPoolExecutor -> song song thuc su vi la subprocess.
    """
    try:
        from PIL import Image
        import io as _io

        # Doc anh bang Pillow -> convert RGB -> ghi PNG vao bytes buffer
        with Image.open(image_path) as img:
            buf = _io.BytesIO()
            img.convert("RGB").save(buf, format="PNG")
            png_bytes = buf.getvalue()

        # Pipe PNG bytes vao stdin cua Tesseract (stdin mode: ten file = "stdin")
        result = subprocess.run(
            ["tesseract", "stdin", "stdout",
             "-l", TESS_LANG, "--psm", TESS_PSM, "--oem", "3"],
            input=png_bytes,
            capture_output=True,
            timeout=60
        )
        raw_text = result.stdout.decode("utf-8", errors="replace").strip()
        if not raw_text:
            safe_print(f"  [OCR-RONG] {image_path.name}")
            return image_path, None
        safe_print(f"  [OCR-OK]   {image_path.name}  ({len(raw_text)} ky tu)")
        return image_path, raw_text
    except subprocess.TimeoutExpired:
        safe_print(f"  [OCR-TIMEOUT] {image_path.name}")
        return image_path, None
    except Exception as e:
        safe_print(f"  [OCR-FAIL]  {image_path.name}: {e}")
        return image_path, None


def ocr_all_parallel(images: list) -> dict:
    """
    OCR tat ca anh song song. Tra ve dict {path: raw_text}.
    """
    workers = min(MAX_OCR_WORKERS, len(images))
    print(f"\n--- Stage 1: Tesseract OCR ({len(images)} anh, {workers} process song song) ---")

    results = {}
    with ThreadPoolExecutor(max_workers=workers) as executor:
        future_to_path = {executor.submit(ocr_one, img): img for img in images}
        done = 0
        for future in as_completed(future_to_path):
            img_path = future_to_path[future]
            _, raw_text = future.result()
            results[img_path] = raw_text
            done += 1
            print(f"  OCR: {done}/{len(images)}", end="\r", flush=True)
    print()
    return results


# ----------------------------------------------------------------------------
# Stage 2: Gom text -> 1 file
# ----------------------------------------------------------------------------
def save_combined_ocr(image_list: list, ocr_map: dict, out_path: Path, subfolder: str = ""):
    """
    Gom tat ca text OCR thanh 1 file duy nhat.
    Format: === CAU i | file: xxx.webp === ... --- END CAU i ---
    """
    raw_dir = OCR_RAW_DIR / subfolder if subfolder else OCR_RAW_DIR
    raw_dir.mkdir(parents=True, exist_ok=True)

    lines = []
    for i, img in enumerate(image_list, start=1):
        text = ocr_map.get(img) or "[OCR_EMPTY]"
        lines.append(f"=== CAU {i} | file: {img.name} ===")
        lines.append(text)
        lines.append(f"--- END CAU {i} ---")
        lines.append("")

    combined = "\n".join(lines)
    out_path.write_text(combined, encoding="utf-8")
    return combined


# ----------------------------------------------------------------------------
# Stage 3: Groq batch (nhieu cau / call)
# ----------------------------------------------------------------------------
def clean_json_array(text: str) -> str:
    """Trich xuat JSON array dau tien tim thay trong text."""
    start = text.find("[")
    if start == -1:
        # Thu tim { -> wrap thanh array
        obj_start = text.find("{")
        if obj_start == -1:
            return text.strip()
        # Tim chuoi JSON object, wrap lai
        return "[" + text[obj_start:].strip() + "]"

    depth, in_str, escape = 0, False, False
    for i, ch in enumerate(text[start:], start=start):
        if escape:
            escape = False
            continue
        if ch == "\\" and in_str:
            escape = True
            continue
        if ch == '"':
            in_str = not in_str
            continue
        if not in_str:
            if ch == "[":
                depth += 1
            elif ch == "]":
                depth -= 1
                if depth == 0:
                    return text[start : i + 1]
    return text.strip()


def parse_json_array(text: str) -> list | None:
    """Parse JSON array, thu sua loi neu can."""
    cleaned = clean_json_array(text)
    try:
        result = json.loads(cleaned)
        if isinstance(result, dict):
            return [result]
        return result
    except json.JSONDecodeError:
        if HAS_JSON_REPAIR:
            try:
                result = json.loads(repair_json(cleaned))
                if isinstance(result, dict):
                    return [result]
                return result
            except Exception:
                pass
    return None


def call_groq_batch(batch_items: list) -> list:
    """
    Goi Groq voi 1 batch gom nhieu cau.
    batch_items: [(idx, image_path, raw_text), ...]
    Tra ve list dict JSON.
    """
    # Tao batch text
    parts = []
    for idx, img_path, raw_text in batch_items:
        parts.append(f"=== CAU {idx} | file: {img_path.name} ===")
        parts.append(raw_text or "[OCR_EMPTY - khong doc duoc text]")
        parts.append(f"--- END CAU {idx} ---")
    batch_text = "\n".join(parts)

    prompt = BATCH_PROMPT_TEMPLATE.format(
        n=len(batch_items),
        idx="{idx}",  # placeholder trong template string (khong thay the)
        fname="{fname}",
        batch_text=batch_text,
    )
    # Fix: thay the lai dung format
    prompt = BATCH_PROMPT_TEMPLATE \
        .replace("{n}", str(len(batch_items))) \
        .replace("{idx}", "i") \
        .replace("{fname}", "file.webp") \
        .replace("{batch_text}", batch_text)

    max_attempts = len(GROQ_API_KEYS) * MAX_JSON_RETRIES
    last_error = None

    for attempt in range(1, max_attempts + 1):
        key = next_key()
        try:
            client = get_groq_client(key)
            response = client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0,
                max_tokens=4096,
            )
            raw_ai = response.choices[0].message.content

            parsed = parse_json_array(raw_ai)
            if parsed and len(parsed) >= 1:
                # Gap vao _source_file neu thieu
                for i, (_, img_path, _) in enumerate(batch_items):
                    if i < len(parsed):
                        if "_source_file" not in parsed[i]:
                            parsed[i]["_source_file"] = img_path.name
                safe_print(f"  [Groq-OK] batch {len(batch_items)} cau")
                return parsed
            else:
                raise ValueError(f"Groq tra ve array rong hoac parse that bai. Raw: {raw_ai[:200]}")

        except Exception as e:
            last_error = e
            err_msg = str(e).lower()
            safe_print(f"  [Groq-ERR lần {attempt}/{max_attempts}] {e}")
            if "rate_limit" in err_msg or "429" in err_msg:
                time.sleep(5)
            elif "503" in err_msg or "overloaded" in err_msg:
                time.sleep(3)
            else:
                time.sleep(1)

    # That bai toan bo batch -> tra ve error dict cho tung cau
    return [
        {"id": img_path.stem, "error": "groq_batch_failed",
         "raw_response": str(last_error), "_source_file": img_path.name}
        for _, img_path, _ in batch_items
    ]


def groq_all_batched(image_list: list, ocr_map: dict) -> list:
    """
    Chia anh thanh cac batch, goi Groq song song tren cac batch.
    Tra ve list dict JSON theo thu tu image_list.
    """
    # Tao list (idx, path, text) theo thu tu
    indexed = [(i + 1, img, ocr_map.get(img)) for i, img in enumerate(image_list)]
    # Chia batch
    batches = [indexed[i:i + BATCH_SIZE] for i in range(0, len(indexed), BATCH_SIZE)]

    print(f"\n--- Stage 3: Groq Batch ({len(indexed)} cau / {BATCH_SIZE} cau per batch = {len(batches)} batch, {MAX_GROQ_WORKERS} song song) ---")

    batch_results: dict = {}  # batch_idx -> list of dicts

    with ThreadPoolExecutor(max_workers=MAX_GROQ_WORKERS) as executor:
        future_to_bidx = {
            executor.submit(call_groq_batch, batch): bidx
            for bidx, batch in enumerate(batches)
        }
        done = 0
        for future in as_completed(future_to_bidx):
            bidx = future_to_bidx[future]
            try:
                batch_results[bidx] = future.result()
            except Exception as e:
                batch = batches[bidx]
                batch_results[bidx] = [
                    {"id": img_path.stem, "error": "exception",
                     "raw_response": str(e), "_source_file": img_path.name}
                    for _, img_path, _ in batch
                ]
            done += 1
            print(f"  Batch: {done}/{len(batches)} xong", flush=True)

    # Ghep ket qua theo thu tu goc
    flat_results = []
    for bidx in range(len(batches)):
        flat_results.extend(batch_results.get(bidx, []))

    # Ket qua co the nhieu/it hon so voi so anh (Groq tra sai so luong)
    # -> map lai theo _source_file
    file_to_result = {}
    for r in flat_results:
        sf = r.get("_source_file")
        if sf:
            file_to_result[sf] = r

    final = []
    for img in image_list:
        if img.name in file_to_result:
            final.append(file_to_result[img.name])
        else:
            # Tim theo idx
            final.append({
                "id": img.stem, "error": "not_in_groq_response",
                "_source_file": img.name
            })
    return final


# ----------------------------------------------------------------------------
# Natural sort
# ----------------------------------------------------------------------------
def natural_sort_key(p: Path):
    return [int(t) if t.isdigit() else t.lower() for t in re.split(r"(\d+)", p.name)]


# ----------------------------------------------------------------------------
# Xuat DOCX
# ----------------------------------------------------------------------------
def build_docx(questions: list, out_path: Path):
    doc = Document()
    for i, q in enumerate(questions):
        q_num = i + 1
        if q.get("error"):
            p   = doc.add_paragraph()
            run = p.add_run(f"Câu {q_num}: [LOI -- {q.get('error')} -- {q.get('_source_file')}]")
            run.bold = True
            doc.add_paragraph()
            continue

        # Phan than cau hoi
        p = doc.add_paragraph()
        run_num = p.add_run(f"Câu {q_num}: ")
        run_num.bold = True
        question_text = str(q.get("question", "")).replace("\\n", "\n").strip()
        p.add_run(question_text)

        # Cac lua chon - dam bao thu tu A B C D
        options = q.get("options", {})
        for letter in sorted(options.keys()):
            text = options[letter]
            if not text:
                continue
            opt_p = doc.add_paragraph()
            run_letter = opt_p.add_run(f"{letter}. ")
            run_letter.bold = True
            opt_p.add_run(str(text))

        # Dong cach
        doc.add_paragraph()

    doc.save(out_path)


# ----------------------------------------------------------------------------
# Pipeline chinh cho 1 folder/set anh
# ----------------------------------------------------------------------------
def run_pipeline(images: list, out_docx: Path, raw_json_path: Path,
                 combined_ocr_path: Path, label: str = ""):
    print(f"\n{'='*55}")
    print(f"  Set    : {label}  ({len(images)} anh)")
    print(f"  DOCX   : {out_docx.name}")
    print(f"  JSON   : {raw_json_path.name}")
    print(f"{'='*55}")

    # Stage 1: OCR song song
    ocr_map = ocr_all_parallel(images)

    # Stage 2: Gom text
    print(f"\n--- Stage 2: Gom OCR text -> {combined_ocr_path.name} ---")
    save_combined_ocr(images, ocr_map, combined_ocr_path, subfolder=label)
    n_ocr_ok = sum(1 for v in ocr_map.values() if v)
    n_ocr_fail = len(images) - n_ocr_ok
    print(f"  OCR thanh cong: {n_ocr_ok} | That bai/rong: {n_ocr_fail}")

    # Stage 3: Groq batch
    results = groq_all_batched(images, ocr_map)

    # Luu JSON tong hop de debug
    raw_json_path.write_text(
        json.dumps(results, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )

    # Build DOCX
    build_docx(results, out_docx)

    n_err = sum(1 for r in results if r.get("error"))
    print(f"\n  -> DOCX : {out_docx}")
    print(f"  -> JSON : {raw_json_path}")
    print(f"  -> OCR  : {combined_ocr_path}")
    print(f"  -> Tong ket: {len(results)} cau | {n_err} loi")
    return results


# ----------------------------------------------------------------------------
# Main
# ----------------------------------------------------------------------------
def main():
    INPUT_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)
    OCR_RAW_DIR.mkdir(exist_ok=True)

    print(f"\nOCR    : Tesseract ({TESS_LANG}) x{MAX_OCR_WORKERS} song song")
    print(f"AI     : Groq/{GROQ_MODEL}, batch {BATCH_SIZE} cau/call, x{MAX_GROQ_WORKERS} song song")
    print(f"Keys   : {len(GROQ_API_KEYS)} API key(s)")

    subfolders = sorted(
        [p for p in INPUT_DIR.iterdir() if p.is_dir()],
        key=natural_sort_key,
    )

    if subfolders:
        print(f"\nTim thay {len(subfolders)} folder")
        for folder in subfolders:
            images = sorted(
                [p for p in folder.iterdir() if p.suffix.lower() in SUPPORTED_EXT],
                key=natural_sort_key,
            )
            if not images:
                print(f"  [bo qua] Khong co anh trong: {folder.name}")
                continue

            out_docx          = OUTPUT_DIR / (folder.name + ".docx")
            raw_json_path     = OUTPUT_DIR / (folder.name + "_raw.json")
            combined_ocr_path = OCR_RAW_DIR / (folder.name + "_all_ocr.txt")

            run_pipeline(images, out_docx, raw_json_path,
                         combined_ocr_path, label=folder.name)
    else:
        images = sorted(
            [p for p in INPUT_DIR.iterdir() if p.suffix.lower() in SUPPORTED_EXT],
            key=natural_sort_key,
        )
        if not images:
            print(f"\nKhong tim thay anh hoac folder nao trong {INPUT_DIR}.")
            print("Hay bo anh vao input_images/ hoac tao subfolder ben trong.")
            sys.exit(0)

        print(f"\nTim thay {len(images)} anh (flat mode).")
        out_docx          = OUTPUT_DIR / "questions.docx"
        raw_json_path     = OUTPUT_DIR / "all_raw.json"
        combined_ocr_path = OCR_RAW_DIR / "all_ocr.txt"

        run_pipeline(images, out_docx, raw_json_path,
                     combined_ocr_path, label="flat")

    print("\nHoan thanh!")


if __name__ == "__main__":
    main()
